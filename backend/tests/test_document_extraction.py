"""Unit tests for in-process document text extraction."""
import io

import pytest

from app.services import document_extraction as de


def _docx_bytes(paragraphs, table_rows=None):
    import docx

    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_is_text_extractable():
    assert de.is_text_extractable("a.docx")
    assert de.is_text_extractable("a.PDF")
    assert de.is_text_extractable("a.txt")
    assert not de.is_text_extractable("a.doc")
    assert not de.is_text_extractable("a.png")


def test_extract_docx_paragraphs_and_tables():
    content = _docx_bytes(
        ["Lessor: Acme Properties", "Base Rent: $5,000"],
        table_rows=[["Year", "Rent"], ["2024", "$5,000"]],
    )
    text = de.extract_text(content, "lease.docx")
    assert "Acme Properties" in text
    assert "Base Rent: $5,000" in text
    # Table cells are included joined by ' | '.
    assert "Year | Rent" in text


def test_extract_txt():
    text = de.extract_text(b"Hello lease world", "notes.txt")
    assert text == "Hello lease world"


def test_extract_doc_raises_unsupported():
    with pytest.raises(de.UnsupportedDocumentError):
        de.extract_text(b"\xd0\xcf\x11\xe0", "old.doc")


def test_extract_unknown_extension_raises():
    with pytest.raises(de.UnsupportedDocumentError):
        de.extract_text(b"data", "image.png")
