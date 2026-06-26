"""Tests for lease document search (keyword + semantic)."""
import uuid

import pytest

from app.auth.jwt_handler import create_access_token
from app.auth.password import hash_password
from app.models.attachment import Attachment
from app.models.lease import Lease
from app.models.lease_document_chunk import LeaseDocumentChunk
from app.models.organization import Organization
from app.models.user import User
from app.services import ai_service, document_search_service


async def _make_org_user(db_session, email: str):
    org = Organization(name="DS Org", slug=f"ds-{email[:6]}", plan="pro")
    db_session.add(org)
    await db_session.commit()
    await db_session.refresh(org)
    user = User(
        email=email,
        display_name="U",
        password_hash=hash_password("x"),
        auth_provider="internal",
        role="admin",
        is_active=True,
        organization_id=org.id,
    )
    db_session.add(user)
    await db_session.commit()
    await db_session.refresh(user)
    token = create_access_token({"sub": str(user.id), "role": user.role})
    return org, user, {"Authorization": "Bearer " + token}


async def _make_lease(db_session, org, name="Acme HQ Lease"):
    lease = Lease(lease_name=name, organization_id=org.id, expiration_year=2030)
    db_session.add(lease)
    await db_session.commit()
    await db_session.refresh(lease)
    return lease


# ── Unit: chunking + similarity ───────────────────────────────────────────────

def test_chunk_text_splits_and_overlaps():
    text = "word " * 1000  # 5000 chars
    chunks = document_search_service.chunk_text(text)
    assert len(chunks) > 1
    assert all(c.strip() for c in chunks)


def test_chunk_text_empty():
    assert document_search_service.chunk_text("") == []
    assert document_search_service.chunk_text("   ") == []


def test_cosine_similarity():
    assert document_search_service._cosine([1, 0], [1, 0]) == pytest.approx(1.0)
    assert document_search_service._cosine([1, 0], [0, 1]) == pytest.approx(0.0)
    assert document_search_service._cosine([], [1, 2]) == 0.0


# ── Keyword fallback (no Gemini) ──────────────────────────────────────────────

@pytest.mark.asyncio
async def test_keyword_search_when_ai_unconfigured(db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, _ = await _make_org_user(db_session, "kw@test.com")
    lease = await _make_lease(db_session, org)

    db_session.add(
        LeaseDocumentChunk(
            organization_id=org.id,
            lease_id=lease.id,
            attachment_id=None,
            source_filename="lease.pdf",
            chunk_index=0,
            content="The tenant shall pay base rent of $5,000 per month for the premises.",
            embedding=None,
        )
    )
    await db_session.commit()

    results = await document_search_service.search_documents(
        db_session, organization_id=org.id, query="base rent"
    )
    assert len(results) == 1
    assert results[0]["lease_id"] == str(lease.id)
    assert results[0]["match_type"] == "keyword"
    assert "base rent" in results[0]["snippet"].lower()


@pytest.mark.asyncio
async def test_keyword_search_is_org_scoped(db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org_a, _, _ = await _make_org_user(db_session, "orga@test.com")
    org_b, _, _ = await _make_org_user(db_session, "orgb@test.com")
    lease_a = await _make_lease(db_session, org_a, "A Lease")
    lease_b = await _make_lease(db_session, org_b, "B Lease")
    for org, lease in ((org_a, lease_a), (org_b, lease_b)):
        db_session.add(
            LeaseDocumentChunk(
                organization_id=org.id,
                lease_id=lease.id,
                source_filename="x.pdf",
                chunk_index=0,
                content="confidential indemnification clause",
            )
        )
    await db_session.commit()

    results = await document_search_service.search_documents(
        db_session, organization_id=org_a.id, query="indemnification"
    )
    assert len(results) == 1
    assert results[0]["lease_id"] == str(lease_a.id)


# ── Semantic ranking (mocked embeddings) ──────────────────────────────────────

@pytest.mark.asyncio
async def test_semantic_search_ranks_by_cosine(db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: True)
    org, user, _ = await _make_org_user(db_session, "sem@test.com")
    lease1 = await _make_lease(db_session, org, "Relevant Lease")
    lease2 = await _make_lease(db_session, org, "Irrelevant Lease")

    db_session.add_all([
        LeaseDocumentChunk(
            organization_id=org.id, lease_id=lease1.id, source_filename="a.pdf",
            chunk_index=0, content="termination rights", embedding=[1.0, 0.0, 0.0],
        ),
        LeaseDocumentChunk(
            organization_id=org.id, lease_id=lease2.id, source_filename="b.pdf",
            chunk_index=0, content="parking provisions", embedding=[0.0, 1.0, 0.0],
        ),
    ])
    await db_session.commit()

    async def fake_embed(texts):
        # Query embedding aligned with lease1's chunk vector.
        return [[1.0, 0.0, 0.0] for _ in texts]

    monkeypatch.setattr(ai_service, "embed_texts", fake_embed)

    results = await document_search_service.search_documents(
        db_session, organization_id=org.id, query="how can the lease be terminated"
    )
    assert results
    assert results[0]["lease_id"] == str(lease1.id)
    assert results[0]["match_type"] == "semantic"
    assert results[0]["score"] == pytest.approx(1.0, abs=1e-3)


# ── Endpoint ──────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_document_search_endpoint(client, db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, headers = await _make_org_user(db_session, "ep@test.com")
    lease = await _make_lease(db_session, org)
    db_session.add(
        LeaseDocumentChunk(
            organization_id=org.id, lease_id=lease.id, source_filename="lease.pdf",
            chunk_index=0, content="renewal option exercisable with 180 days notice",
        )
    )
    await db_session.commit()

    resp = await client.post(
        f"/api/v1/leases/{lease.id}/document-search",
        headers=headers,
        json={"query": "renewal option"},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["matches"]
    assert body["matches"][0]["lease_id"] == str(lease.id)


@pytest.mark.asyncio
async def test_index_attachment_keyword_only(db_session, monkeypatch):
    """index_attachment stores chunks even when AI is off (keyword-only)."""
    import io

    import docx

    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, _ = await _make_org_user(db_session, "idx@test.com")
    lease = await _make_lease(db_session, org)

    document = docx.Document()
    document.add_paragraph("Assignment and subletting require landlord consent.")
    buf = io.BytesIO()
    document.save(buf)

    attachment = Attachment(
        organization_id=org.id,
        entity_type="lease",
        entity_id=lease.id,
        original_filename="lease.docx",
        stored_filename=f"{uuid.uuid4()}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=buf.tell(),
        uploaded_by=user.email,
    )
    db_session.add(attachment)
    await db_session.commit()
    await db_session.refresh(attachment)

    count = await document_search_service.index_attachment(
        db_session, lease=lease, attachment=attachment, content=buf.getvalue()
    )
    assert count == 1

    results = await document_search_service.search_documents(
        db_session, organization_id=org.id, query="subletting"
    )
    assert results and results[0]["lease_id"] == str(lease.id)


def test_snippet_returns_full_paragraph_single_newline():
    """Word-style text (paragraphs joined by single newlines) returns the whole
    matched paragraph, not a truncated window."""
    content = (
        "Article 1. Premises. The premises are leased as-is.\n"
        "Article 2. Base Rent. Tenant shall pay base rent of $28,500 per month "
        "payable monthly in advance without demand, setoff, or deduction.\n"
        "Article 3. Term. The term is sixty months."
    )
    snip = document_search_service._snippet(content, "base rent")
    assert snip.startswith("Article 2. Base Rent.")
    assert "without demand, setoff, or deduction." in snip
    # Must not bleed into neighbouring paragraphs.
    assert "Article 1." not in snip
    assert "Article 3." not in snip


def test_snippet_returns_full_paragraph_blank_line_delimited():
    """PDF-style text (paragraphs separated by blank lines, wrapped with single
    newlines) returns the full paragraph with line breaks collapsed."""
    content = (
        "Page intro text.\n\n"
        "5. Renewal and Notice\n"
        "Tenant shall have one option to renew\n"
        "for five years by delivering written\n"
        "notice ninety days prior to expiration.\n\n"
        "6. Permitted Use\nGeneral office use only."
    )
    snip = document_search_service._snippet(content, "ninety days")
    assert "5. Renewal and Notice" in snip
    assert "ninety days prior to expiration." in snip
    assert "\n" not in snip  # internal line breaks collapsed
    assert "Permitted Use" not in snip


def test_snippet_bounds_overlong_paragraph():
    content = "x" * 4000 + " base rent clause " + "y" * 4000
    snip = document_search_service._snippet(content, "base rent", max_chars=120)
    assert "base rent" in snip
    assert len(snip) <= 122  # max_chars plus the two ellipsis characters
    assert snip.startswith("…") and snip.endswith("…")


# ── Multiple hits per document (single-lease search) ──────────────────────────

@pytest.mark.asyncio
async def test_single_lease_search_returns_multiple_hits(db_session, monkeypatch):
    """A single-lease search surfaces every matching chunk (no per-lease collapse)
    so the preview pane can navigate between occurrences."""
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, _user, _ = await _make_org_user(db_session, "multi@test.com")
    lease = await _make_lease(db_session, org)
    for idx in range(3):
        db_session.add(
            LeaseDocumentChunk(
                organization_id=org.id, lease_id=lease.id, source_filename="lease.pdf",
                chunk_index=idx, content=f"clause {idx} mentions base rent details",
            )
        )
    await db_session.commit()

    results = await document_search_service.search_documents(
        db_session, organization_id=org.id, query="base rent", lease_id=lease.id
    )
    assert len(results) == 3
    assert {r["chunk_index"] for r in results} == {0, 1, 2}


@pytest.mark.asyncio
async def test_portfolio_search_collapses_per_lease(db_session, monkeypatch):
    """A portfolio-wide search still collapses to one best match per lease."""
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, _user, _ = await _make_org_user(db_session, "collapse@test.com")
    lease = await _make_lease(db_session, org)
    for idx in range(3):
        db_session.add(
            LeaseDocumentChunk(
                organization_id=org.id, lease_id=lease.id, source_filename="lease.pdf",
                chunk_index=idx, content=f"clause {idx} mentions base rent details",
            )
        )
    await db_session.commit()

    results = await document_search_service.search_documents(
        db_session, organization_id=org.id, query="base rent"
    )
    assert len(results) == 1


# ── Document text preview ─────────────────────────────────────────────────────

async def _make_lease_docx_attachment(db_session, org, user, paragraphs):
    import io
    from pathlib import Path

    import docx

    from app.config import settings

    lease = await _make_lease(db_session, org)
    document = docx.Document()
    for para in paragraphs:
        document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    stored = f"{uuid.uuid4()}.docx"
    attachment = Attachment(
        organization_id=org.id,
        entity_type="lease",
        entity_id=lease.id,
        original_filename="lease.docx",
        stored_filename=stored,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=buf.tell(),
        uploaded_by=user.email,
    )
    db_session.add(attachment)
    await db_session.commit()
    await db_session.refresh(attachment)

    dest_dir = Path(settings.UPLOAD_DIR) / "lease"
    dest_dir.mkdir(parents=True, exist_ok=True)
    (dest_dir / stored).write_bytes(buf.getvalue())
    return lease, attachment, dest_dir / stored


@pytest.mark.asyncio
async def test_get_document_text_service(db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, _ = await _make_org_user(db_session, "txt@test.com")
    lease, attachment, path = await _make_lease_docx_attachment(
        db_session, org, user, ["Base rent is $10,000.", "Term is sixty months."]
    )
    try:
        result = await document_search_service.get_document_text(
            db_session, lease=lease, attachment_id=attachment.id
        )
    finally:
        path.unlink(missing_ok=True)
    assert result is not None
    assert result["extractable"] is True
    assert "Base rent is $10,000." in result["text"]
    assert "Term is sixty months." in result["text"]


@pytest.mark.asyncio
async def test_get_document_text_wrong_lease_returns_none(db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, _ = await _make_org_user(db_session, "wrong@test.com")
    lease, attachment, path = await _make_lease_docx_attachment(
        db_session, org, user, ["Base rent is $10,000."]
    )
    other_lease = await _make_lease(db_session, org, name="Other Lease")
    try:
        result = await document_search_service.get_document_text(
            db_session, lease=other_lease, attachment_id=attachment.id
        )
    finally:
        path.unlink(missing_ok=True)
    assert result is None


@pytest.mark.asyncio
async def test_document_text_endpoint(client, db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, user, headers = await _make_org_user(db_session, "txtep@test.com")
    lease, attachment, path = await _make_lease_docx_attachment(
        db_session, org, user, ["Commencement date is January 1.", "Base rent clause."]
    )
    try:
        resp = await client.get(
            f"/api/v1/leases/{lease.id}/documents/{attachment.id}/text",
            headers=headers,
        )
    finally:
        path.unlink(missing_ok=True)
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["extractable"] is True
    assert "Commencement date is January 1." in body["text"]
    assert body["source_filename"] == "lease.docx"


@pytest.mark.asyncio
async def test_document_text_endpoint_unknown_attachment(client, db_session, monkeypatch):
    monkeypatch.setattr(ai_service, "is_configured", lambda: False)
    org, _user, headers = await _make_org_user(db_session, "missing@test.com")
    lease = await _make_lease(db_session, org)
    resp = await client.get(
        f"/api/v1/leases/{lease.id}/documents/{uuid.uuid4()}/text",
        headers=headers,
    )
    assert resp.status_code == 404
