"""Plain-text extraction from uploaded documents.

Shared by:

* the AI lease parser (``ai_service``/``ai`` router) — Gemini can read PDF and
  image bytes inline but **cannot** parse Office binary formats, so Word
  documents are converted to text here before being handed to the model; and
* the lease document search index (``document_search_service``) — every text
  document attached to a lease has its content extracted and indexed.

Supported inputs: ``.pdf`` (via :mod:`pypdf`), ``.docx`` (via
:mod:`python-docx`) and ``.txt``. Legacy binary ``.doc`` is intentionally *not*
supported (it needs an external converter); callers should surface a clear
"convert to .docx or PDF" message instead.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Extensions we can turn into plain text in-process.
TEXT_EXTRACTABLE_EXTENSIONS = frozenset({".pdf", ".docx", ".txt"})

# Hard cap on extracted characters so a pathological document can't blow up the
# prompt or the database row.
MAX_EXTRACTED_CHARS = 500_000


class DocumentExtractionError(Exception):
    """Base class for extraction failures."""


class UnsupportedDocumentError(DocumentExtractionError):
    """Raised when the file type cannot be extracted to text in-process."""


def is_text_extractable(filename: str) -> bool:
    """Return whether :func:`extract_text` can handle ``filename``'s type."""
    return Path(filename).suffix.lower() in TEXT_EXTRACTABLE_EXTENSIONS


def extract_text(content: bytes, filename: str, *, max_chars: int | None = None) -> str:
    """Extract plain text from ``content`` based on ``filename``'s extension.

    ``max_chars`` overrides the default :data:`MAX_EXTRACTED_CHARS` ceiling —
    the AI document pipeline raises it so a large document can be segmented and
    processed in full instead of being silently truncated.

    Raises :class:`UnsupportedDocumentError` for unsupported types (including the
    legacy binary ``.doc`` format) and :class:`DocumentExtractionError` when a
    supported document cannot be parsed.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(content)
    elif ext == ".docx":
        text = _extract_docx(content)
    elif ext == ".txt":
        text = _extract_txt(content)
    elif ext == ".doc":
        raise UnsupportedDocumentError(
            "Legacy .doc files are not supported. Please convert the document to "
            ".docx or PDF and try again."
        )
    else:
        raise UnsupportedDocumentError(
            f"Cannot extract text from '{ext}' files. Supported: .pdf, .docx, .txt."
        )
    limit = MAX_EXTRACTED_CHARS if max_chars is None else max(1, max_chars)
    return text[:limit].strip()


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise DocumentExtractionError("pypdf is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read PDF: {exc}") from exc
    return "\n\n".join(p for p in pages if p.strip())


def split_pdf(content: bytes, *, max_bytes: int, max_parts: int) -> list[bytes]:
    """Split a PDF into page-range PDFs that each fit within ``max_bytes``.

    Used for documents (typically scanned/image-only PDFs) that are too large to
    hand to the model in one piece and have no text layer to extract. Pages are
    grouped in reading order so every page appears in exactly one part and no
    content is dropped. At most ``max_parts`` parts are returned; anything beyond
    that is left out rather than producing an unbounded number of model calls.

    Raises :class:`DocumentExtractionError` when the PDF cannot be read or when a
    single page still exceeds ``max_bytes``.
    """
    if max_bytes <= 0 or max_parts <= 0:
        raise DocumentExtractionError("Invalid PDF split bounds")
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise DocumentExtractionError("pypdf is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(content))
        page_count = len(reader.pages)
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read PDF: {exc}") from exc
    if page_count == 0:
        raise DocumentExtractionError("The PDF did not contain any pages")

    def _serialize(start: int, end: int) -> bytes:
        writer = PdfWriter()
        for index in range(start, end):
            writer.add_page(reader.pages[index])
        buffer = io.BytesIO()
        try:
            writer.write(buffer)
        except Exception as exc:
            raise DocumentExtractionError(f"Could not split PDF: {exc}") from exc
        return buffer.getvalue()

    # Start from a size-proportional page group and shrink it if pypdf's output
    # (which re-embeds shared resources per part) overshoots the byte ceiling.
    pages_per_part = max(1, int(page_count * max_bytes / max(len(content), 1) * 0.9))
    parts: list[bytes] = []
    start = 0
    while start < page_count and len(parts) < max_parts:
        end = min(start + pages_per_part, page_count)
        blob = _serialize(start, end)
        while len(blob) > max_bytes and end - start > 1:
            end = start + max(1, (end - start) // 2)
            blob = _serialize(start, end)
        if len(blob) > max_bytes:
            raise DocumentExtractionError(
                "A single page of this PDF is too large for AI processing."
            )
        parts.append(blob)
        pages_per_part = max(1, end - start)
        start = end
    return parts


def _extract_docx(content: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise DocumentExtractionError("python-docx is not installed") from exc
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError(f"Could not read Word document: {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text — leases frequently put rent/term tables in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")
