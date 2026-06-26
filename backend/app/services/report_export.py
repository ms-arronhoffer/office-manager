"""Render an AI briefing (Markdown) to HTML, PDF and DOCX.

The AI operations briefing is produced as Markdown by ``ai_service``. This
service converts that Markdown into:

* **HTML** — for in-app display and HTML email bodies (via :mod:`markdown`).
* **PDF** — via ReportLab (already a project dependency).
* **DOCX** — via python-docx.

Keeping all three conversions here means the API endpoint, the scheduled email
job, and the download buttons all share one implementation.
"""
from __future__ import annotations

import io
import re

EXPORT_FORMATS = ("pdf", "docx")

# Markdown line patterns. The captured text group begins with ``\S`` so it can
# never overlap the preceding ``\s+``; this keeps matching linear and avoids
# polynomial backtracking (ReDoS) on adversarial whitespace-heavy input.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(\S.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(\S.*)$")
_ORDERED_RE = re.compile(r"^\d+[.)]\s+(\S.*)$")


def markdown_to_html(md_text: str) -> str:
    """Convert Markdown to an HTML fragment.

    Falls back to a ``<pre>``-wrapped, HTML-escaped block if the optional
    ``markdown`` dependency is unavailable, so callers always get safe HTML.
    """
    try:
        import markdown as _markdown
    except ImportError:  # pragma: no cover - dependency guard
        from html import escape

        return f"<pre>{escape(md_text or '')}</pre>"
    return _markdown.markdown(
        md_text or "",
        extensions=["extra", "sane_lists", "nl2br"],
    )


def markdown_to_email_html(md_text: str, *, title: str | None = None) -> str:
    """Wrap the rendered briefing in a minimal, email-friendly HTML document."""
    body = markdown_to_html(md_text)
    heading = f"<h1>{title}</h1>" if title else ""
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:Arial,Helvetica,sans-serif;color:#16191f;"
        "line-height:1.5;max-width:680px;margin:0 auto;padding:16px}"
        "h1,h2,h3{color:#0f1b2d}table{border-collapse:collapse}"
        "td,th{border:1px solid #ccc;padding:4px 8px}</style></head>"
        f"<body>{heading}{body}</body></html>"
    )


# ── DOCX ──────────────────────────────────────────────────────────────────────

def markdown_to_docx(md_text: str, *, title: str | None = None) -> bytes:
    """Render Markdown to a .docx document and return its bytes.

    Supports the subset of Markdown the briefing uses: ATX headings (``#``..),
    bullet/numbered lists, and paragraphs (with inline ``**bold**`` stripped to
    plain text for simplicity).
    """
    import docx

    document = docx.Document()
    if title:
        document.add_heading(title, level=0)

    for raw_line in (md_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(_strip_inline(heading_match.group(2)), level=level)
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            document.add_paragraph(_strip_inline(bullet_match.group(1)), style="List Bullet")
            continue

        ordered_match = _ORDERED_RE.match(stripped)
        if ordered_match:
            document.add_paragraph(_strip_inline(ordered_match.group(1)), style="List Number")
            continue

        document.add_paragraph(_strip_inline(stripped))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def markdown_to_pdf(md_text: str, *, title: str | None = None) -> bytes:
    """Render Markdown to a PDF and return its bytes (via ReportLab)."""
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "BriefingBody", parent=styles["BodyText"], alignment=TA_LEFT, spaceAfter=8
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title=title or "Briefing",
    )

    flow = []
    if title:
        flow.append(Paragraph(_inline_to_pdf(title), styles["Title"]))
        flow.append(Spacer(1, 12))

    pending_bullets: list[str] = []

    def flush_bullets():
        if pending_bullets:
            flow.append(
                ListFlowable(
                    [ListItem(Paragraph(_inline_to_pdf(b), body_style)) for b in pending_bullets],
                    bulletType="bullet",
                )
            )
            pending_bullets.clear()

    for raw_line in (md_text or "").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            flush_bullets()
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            flush_bullets()
            level = len(heading_match.group(1))
            style = styles["Heading1"] if level <= 1 else styles["Heading2"] if level == 2 else styles["Heading3"]
            flow.append(Paragraph(_inline_to_pdf(heading_match.group(2)), style))
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            pending_bullets.append(bullet_match.group(1))
            continue

        ordered_match = _ORDERED_RE.match(stripped)
        if ordered_match:
            pending_bullets.append(ordered_match.group(1))
            continue

        flush_bullets()
        flow.append(Paragraph(_inline_to_pdf(stripped), body_style))

    flush_bullets()
    if not flow:
        flow.append(Paragraph("(No content)", body_style))

    doc.build(flow)
    return buf.getvalue()


# ── Inline helpers ────────────────────────────────────────────────────────────

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
_ITALIC_RE = re.compile(r"\*(.+?)\*|_(.+?)_")


def _strip_inline(text: str) -> str:
    """Remove Markdown emphasis markers, leaving plain text (for DOCX)."""
    text = _BOLD_RE.sub(lambda m: m.group(1) or m.group(2), text)
    text = _ITALIC_RE.sub(lambda m: m.group(1) or m.group(2), text)
    return text


def _inline_to_pdf(text: str) -> str:
    """Convert simple Markdown emphasis to ReportLab mini-HTML, escaping the rest."""
    from xml.sax.saxutils import escape

    escaped = escape(text)
    escaped = _BOLD_RE.sub(lambda m: f"<b>{m.group(1) or m.group(2)}</b>", escaped)
    escaped = _ITALIC_RE.sub(lambda m: f"<i>{m.group(1) or m.group(2)}</i>", escaped)
    return escaped
